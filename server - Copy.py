import asyncio
import os
from pathlib import Path
from zipfile import ZipFile
from docx import Document

from mcp.server.fastmcp import FastMCP

# Initialize MCP server
mcp = FastMCP("docx-mcp")

# Base directory for documents
DOCS_DIR = Path("./documents")
DOCS_DIR.mkdir(exist_ok=True)

def get_doc_path(filename: str) -> Path:
    if not filename.endswith(".docx"):
        filename += ".docx"
    return DOCS_DIR / filename

# -------------------------------
# TOOLS
# -------------------------------

@mcp.tool()
def create_docx(filename: str, content: str = "") -> str:
    """
    Create a new Word document with optional initial content.
    """
    path = get_doc_path(filename)
    if path.exists():
        return f"❌ File already exists: {path}"

    try:
        doc = Document()
        if content.strip():
            doc.add_paragraph(content)
        else:
            doc.add_paragraph("(new document)")
        doc.save(path)

        # Verify it opens
        _ = Document(path)
        return f"✅ Created document: {path}"
    except Exception as e:
        return f"❌ Failed to create document: {str(e)}"

@mcp.tool()
def read_docx(filename: str) -> str:
    """
    Read all text from a Word document, with validation.
    """
    path = get_doc_path(filename)
    if not path.exists():
        return f"❌ File not found: {path}"

    try:
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text or "(empty document)"
    except Exception as e:
        # Inspect package to see what it really is
        try:
            with ZipFile(path, 'r') as zf:
                files = zf.namelist()
                if "word/document.xml" not in files:
                    return (
                        f"❌ {filename} is an OOXML container but not a Word document.\n"
                        f"Contains: {files[:10]}..."
                    )
                else:
                    return f"❌ Could not read {filename}: {str(e)}"
        except Exception:
            return f"❌ {filename} is not a valid .docx (possibly corrupted)."

@mcp.tool()
def update_docx(filename: str, changes: list[str]) -> str:
    """
    Append new paragraphs to a Word document.
    """
    path = get_doc_path(filename)
    if not path.exists():
        return f"❌ File not found: {path}"

    try:
        doc = Document(path)
        for change in changes:
            doc.add_paragraph(change)
        doc.save(path)
        return f"✅ Updated {filename} with {len(changes)} changes."
    except Exception as e:
        return f"❌ Failed to update {filename}: {str(e)}"

@mcp.tool()
def delete_docx(filename: str) -> str:
    """
    Delete a Word document.
    """
    path = get_doc_path(filename)
    if not path.exists():
        return f"❌ File not found: {path}"

    try:
        os.remove(path)
        return f"🗑️ Deleted: {path}"
    except Exception as e:
        return f"❌ Failed to delete {filename}: {str(e)}"

@mcp.tool()
def convert_to_docx(input_file: str, output_file: str) -> str:
    """
    Convert a .txt file (and similar) into a valid .docx.
    """
    in_path = Path(input_file)
    out_path = get_doc_path(output_file)

    if not in_path.exists():
        return f"❌ Input file not found: {in_path}"

    try:
        # Try reading as plain text
        with open(in_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(out_path)

        return f"✅ Converted {in_path} → {out_path}"
    except Exception as e:
        return f"❌ Could not convert {input_file}: {str(e)}"

# -------------------------------
# MAIN ENTRY
# -------------------------------
if __name__ == "__main__":
    asyncio.run(mcp.run())
